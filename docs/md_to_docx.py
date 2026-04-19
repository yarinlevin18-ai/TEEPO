# -*- coding: utf-8 -*-
"""
Convert TEEPO_SPEC.md (Hebrew, RTL) to a nicely-formatted .docx.

- Headings H1/H2/H3 preserved
- Bullet lists, numbered lists
- Pipe tables
- Bold (**text**), italic (*text*), inline code (`text`)
- RTL direction applied to every paragraph + table cell
- Default font: David (Hebrew-friendly), falls back to Arial
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor

SRC = Path(r"C:\dev\smartdesk\docs\TEEPO_SPEC.md")
DST = Path(r"C:\dev\smartdesk\docs\TEEPO_SPEC.docx")

FONT_HE = "David"  # Hebrew-friendly; widely available on Windows
FONT_LATIN = "Arial"


# ─────────────────────────────────────────────────────────────
# RTL helpers
# ─────────────────────────────────────────────────────────────
def set_rtl(paragraph):
    """Set paragraph direction to RTL."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    bidi.set(qn("w:val"), "1")
    # Right-align by default for Hebrew
    if paragraph.alignment is None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_run_rtl(run):
    """Mark a run as RTL (complex script) so Hebrew renders correctly."""
    rPr = run._r.get_or_add_rPr()
    rtl = rPr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        rPr.append(rtl)
    rtl.set(qn("w:val"), "1")


def set_cell_rtl(cell):
    """Apply RTL to all paragraphs in a table cell."""
    for p in cell.paragraphs:
        set_rtl(p)


def set_font(run, size=None, bold=None, italic=None, color=None, mono=False):
    """Apply font + size + style to a run."""
    font = run.font
    font.name = "Consolas" if mono else FONT_LATIN
    # Set Hebrew/complex-script font explicitly
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Consolas" if mono else FONT_LATIN)
    rFonts.set(qn("w:hAnsi"), "Consolas" if mono else FONT_LATIN)
    rFonts.set(qn("w:cs"), "Consolas" if mono else FONT_HE)
    rFonts.set(qn("w:eastAsia"), FONT_LATIN)
    if size is not None:
        font.size = Pt(size)
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rPr.append(szCs)
        szCs.set(qn("w:val"), str(size * 2))
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if color is not None:
        font.color.rgb = color


# ─────────────────────────────────────────────────────────────
# Inline formatting (bold, italic, code)
# ─────────────────────────────────────────────────────────────
INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))"
)


def add_inline_runs(paragraph, text, base_size=11):
    """Parse inline markdown (bold/italic/code/links) and add runs."""
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=base_size, bold=True)
            set_run_rtl(run)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=base_size, italic=True)
            set_run_rtl(run)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=base_size - 1, mono=True)
            set_run_rtl(run)
        elif part.startswith("[") and "](" in part:
            m = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", part)
            if m:
                run = paragraph.add_run(m.group(1))
                set_font(run, size=base_size, color=RGBColor(0x1F, 0x4E, 0x79))
                run.underline = True
                set_run_rtl(run)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=base_size)
            set_run_rtl(run)


# ─────────────────────────────────────────────────────────────
# Document setup
# ─────────────────────────────────────────────────────────────
def configure_document(doc):
    """Base document-wide RTL + styles."""
    # Normal style defaults
    style = doc.styles["Normal"]
    style.font.name = FONT_LATIN
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_LATIN)
    rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    rFonts.set(qn("w:cs"), FONT_HE)

    # Headings with sensible sizes
    for level, size in [(1, 20), (2, 16), (3, 13)]:
        name = f"Heading {level}"
        if name in [s.name for s in doc.styles]:
            s = doc.styles[name]
            s.font.size = Pt(size)
            s.font.bold = True
            s.font.color.rgb = RGBColor(0x2E, 0x2E, 0x5A)
            rPr2 = s.element.get_or_add_rPr()
            rFonts2 = rPr2.find(qn("w:rFonts"))
            if rFonts2 is None:
                rFonts2 = OxmlElement("w:rFonts")
                rPr2.append(rFonts2)
            rFonts2.set(qn("w:ascii"), FONT_LATIN)
            rFonts2.set(qn("w:hAnsi"), FONT_LATIN)
            rFonts2.set(qn("w:cs"), FONT_HE)


# ─────────────────────────────────────────────────────────────
# Table parsing (markdown pipe tables)
# ─────────────────────────────────────────────────────────────
def is_table_separator(line):
    stripped = line.strip()
    return bool(re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", stripped))


def parse_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    t.autofit = True
    # Make sure the table itself renders RTL
    tblPr = t._tbl.tblPr
    bidi = tblPr.find(qn("w:bidiVisual"))
    if bidi is None:
        bidi = OxmlElement("w:bidiVisual")
        tblPr.append(bidi)
    bidi.set(qn("w:val"), "1")

    # Header row
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        set_rtl(p)
        add_inline_runs(p, h, base_size=11)
        for r in p.runs:
            r.bold = True
        set_cell_rtl(hdr[i])

    # Body rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            if ci >= len(cells):
                break
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            set_rtl(p)
            add_inline_runs(p, val, base_size=10)
            set_cell_rtl(cells[ci])


# ─────────────────────────────────────────────────────────────
# Main parser
# ─────────────────────────────────────────────────────────────
def convert(src_path, dst_path):
    text = src_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    configure_document(doc)

    # Document-level RTL via sectPr
    section = doc.sections[0]
    sectPr = section._sectPr
    bidi = sectPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        sectPr.append(bidi)

    i = 0
    n = len(lines)
    in_code = False
    code_lang = ""

    while i < n:
        line = lines[i]

        # Fenced code blocks
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line.strip()[3:]
                i += 1
                code_lines = []
                while i < n and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                i += 1  # skip closing fence
                p = doc.add_paragraph()
                set_rtl(p)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # code is LTR
                pPr = p._p.get_or_add_pPr()
                bidi_p = pPr.find(qn("w:bidi"))
                if bidi_p is not None:
                    pPr.remove(bidi_p)
                run = p.add_run("\n".join(code_lines))
                set_font(run, size=9, mono=True)
                # Light gray background via shd
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F2F2F2")
                p._p.get_or_add_pPr().append(shd)
                in_code = False
                continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            set_rtl(p)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "888888")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Tables
        if "|" in line and i + 1 < n and is_table_separator(lines[i + 1]):
            header = parse_row(line)
            i += 2  # skip separator
            rows = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(parse_row(lines[i]))
                i += 1
            add_table(doc, header, rows)
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            lvl = min(level, 3)
            p = doc.add_heading(level=lvl)
            set_rtl(p)
            add_inline_runs(p, heading_text, base_size=[0, 20, 16, 13][lvl])
            i += 1
            continue

        # Bullet list
        m = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        if m:
            indent = len(m.group(1))
            text_part = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            set_rtl(p)
            if indent:
                p.paragraph_format.left_indent = Cm(0.5 * (indent // 2 + 1))
            add_inline_runs(p, text_part, base_size=11)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if m:
            text_part = m.group(2)
            p = doc.add_paragraph(style="List Number")
            set_rtl(p)
            add_inline_runs(p, text_part, base_size=11)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph — merge consecutive non-blank lines
        para_lines = [line]
        i += 1
        while i < n:
            nxt = lines[i]
            if (
                not nxt.strip()
                or nxt.strip().startswith("#")
                or nxt.strip().startswith("```")
                or nxt.strip() in ("---", "***", "___")
                or re.match(r"^\s*[-*]\s+", nxt)
                or re.match(r"^\s*\d+\.\s+", nxt)
                or ("|" in nxt and i + 1 < n and is_table_separator(lines[i + 1]))
            ):
                break
            para_lines.append(nxt)
            i += 1

        p = doc.add_paragraph()
        set_rtl(p)
        add_inline_runs(p, " ".join(para_lines), base_size=11)

    doc.save(dst_path)
    print(f"Saved: {dst_path}")


if __name__ == "__main__":
    convert(SRC, DST)
